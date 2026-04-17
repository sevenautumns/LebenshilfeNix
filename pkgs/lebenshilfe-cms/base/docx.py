import pathlib
import zipfile
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_TEMPLATE_PATH = pathlib.Path(__file__).parent / "data" / "briefbogen.dotx"


def _open_dotx(path: pathlib.Path) -> DocumentType:
    """Öffnet eine .dotx-Vorlage als python-docx Document.

    python-docx lehnt .dotx-Dateien wegen des abweichenden Content-Types ab.
    Diese Funktion patcht den Content-Type im Speicher, ohne die Datei zu ändern.
    """
    with zipfile.ZipFile(path) as zin:
        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"wordprocessingml.template.main+xml",
                        b"wordprocessingml.document.main+xml",
                    )
                zout.writestr(item, data)
        buf.seek(0)
    return Document(buf)


class BriefbogenDocument:
    """Erzeugt ein python-docx-Dokument auf Basis des Lebenshilfe-Briefbogens.

    Öffnet die .dotx-Vorlage, leert den Body und stellt Hilfsmethoden
    für häufige Inhaltsblöcke bereit.
    """

    def __init__(self) -> None:
        self._doc = _open_dotx(_TEMPLATE_PATH)
        self._clear_body()

    # ------------------------------------------------------------------ #
    # Interne Hilfsmethoden                                               #
    # ------------------------------------------------------------------ #

    def _clear_body(self) -> None:
        """Entfernt Musterinhalt der Vorlage, behält sectPr."""
        body = self._doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    @staticmethod
    def _set_table_borders(
        table, val: str, sz: str = "4", color: str = "000000"
    ) -> None:
        """Setzt alle Tabellenrahmen auf den angegebenen Wert (z. B. „none" oder „single")."""
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), val)
            if val != "none":
                border.set(qn("w:sz"), sz)
                border.set(qn("w:color"), color)
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # ------------------------------------------------------------------ #
    # Öffentliche API                                                     #
    # ------------------------------------------------------------------ #

    @property
    def doc(self) -> DocumentType:
        """Zugriff auf das rohe python-docx Document für Sonderfälle."""
        return self._doc

    def add_title(self, text: str):
        """Großer, fetter Dokumenttitel (kein Word-Heading-Style erforderlich)."""
        para = self._doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        return para

    def add_paragraph(self, text: str = "", bold: bool = False):
        """Normaler Fließtext-Paragraph."""
        para = self._doc.add_paragraph(text)
        if bold and text:
            para.runs[0].bold = True
        return para

    def add_spacer(self):
        """Leerer Paragraph als Abstandshalter."""
        return self._doc.add_paragraph()

    def add_borderless_table(self, rows: int, cols: int):
        """Tabelle ohne sichtbare Rahmen (für Metadatenblöcke)."""
        table = self._doc.add_table(rows=rows, cols=cols)
        self._set_table_borders(table, "none")
        return table

    def add_grid_table(self, rows: int, cols: int):
        """Tabelle mit sichtbaren Rahmen (für Zahlentabellen)."""
        table = self._doc.add_table(rows=rows, cols=cols)
        self._set_table_borders(table, "single")
        return table

    def to_response(self, filename: str):
        """Gibt einen Django HttpResponse mit Content-Disposition zurück."""
        from django.http import HttpResponse

        buffer = BytesIO()
        self._doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
