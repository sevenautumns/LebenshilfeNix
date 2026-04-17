from django import forms
from django.contrib import admin
from django.contrib.admin import SimpleListFilter
from django.db.models import Count, Q, Sum
from django.shortcuts import redirect
from django.urls import path, reverse
from django.utils.formats import number_format
from django.utils.html import format_html
from unfold.admin import TabularInline
from unfold.contrib.filters.admin import AutocompleteSelectFilter, RangeDateFilter
from unfold.decorators import action, display
from unfold.enums import ActionVariant

from base.admin import (
    BaseModelAdmin,
    ReadOnlyAdminMixin,
    AddressInline,
    PhoneInline,
    EmailInline,
)
from base.mixins import ListSummaryMixin
from base.admin_views import (
    BaseApplyView,
    BaseCalculatorView,
    BaseUnionListView,
    render_label,
)
from base.fields import EuroDecimalField, HourMinuteDurationField
from .models import (
    School,
    SchoolContact,
    SchoolReport,
    NewRequest,
    Student,
    Supervision,
    TandemPairing,
    Request,
)


@admin.register(NewRequest)
class NewRequestAdmin(ReadOnlyAdminMixin, BaseModelAdmin):
    def changelist_view(self, request, extra_context=None):
        return NewRequestListView.as_view(model_admin=self)(request)


_euro_fmt = EuroDecimalField(max_digits=10, decimal_places=2)


class SupervisionCalculatorView(BaseCalculatorView):
    title = "Betreuungsrechner"

    def get_form_class(self):
        from .forms import SupervisionCalculatorOverridesForm

        return SupervisionCalculatorOverridesForm

    def get_source_fields(self, obj: Supervision):
        from base.fields import HourMinuteDurationField as HMField
        from base.models import SchoolDays
        from django.db.models import Q

        is_tandem = TandemPairing.objects.filter(
            Q(supervision_a=obj) | Q(supervision_b=obj)
        ).exists()

        breakdown = SchoolDays.school_days_breakdown(obj.start_date, obj.end_date)

        return [
            ("Schüler:in", str(obj.student)),
            (
                "Kostenträger",
                str(obj.student.payer) if obj.student_id else "—",  # type: ignore[attr-defined]
            ),
            ("Schule", str(obj.school)),
            (
                "Zeitraum",
                f"{obj.start_date.strftime('%d.%m.%Y')} – {obj.end_date.strftime('%d.%m.%Y')}",
            ),
            (
                "Wochenstunden",
                HMField.format_std(obj.weekly_hours),
            ),
            ("Schultage (lt. Kalender)", breakdown["school_days"]),
            ("Urlaubstage (lt. Kalender)", breakdown["vacation_days"]),
            ("Feiertage (lt. Kalender)", breakdown["public_holidays"]),
            ("Gesamt Schultage (rechnerisch)", breakdown["total"]),
            (
                "Art der Betreuung",
                "Tandembetreuung" if is_tandem else "Einzelbetreuung",
            ),
        ]

    def get_primary_results(self, obj: Supervision, result):
        i = result.input
        override_params = self.overrides_to_params(
            {
                "months_override": i.months_override,
                "school_days_override": i.school_days_override,
                "vacation_days_override": i.vacation_days_override,
                "public_holidays_override": i.public_holidays_override,
                "fee_agreement_override": i.fee_agreement_override,
                "use_fee_agreement": True if i.use_fee_agreement else None,
            }
        )

        apply_total_url = self.build_apply_url(
            obj, "admin:pedagogy_supervision_calculator_apply_total", override_params
        )
        apply_installment_url = self.build_apply_url(
            obj,
            "admin:pedagogy_supervision_calculator_apply_installment",
            override_params,
        )

        if result.is_pool_rate:
            total_suffix = "Poolpauschale"
            installment_suffix = "Poolpauschale"
        elif result.is_tandem:
            total_suffix = "nach Tandemabzug"
            installment_suffix = "nach Tandemabzug"
        else:
            total_suffix = "berechnet"
            installment_suffix = "berechnet"

        return [
            {
                "label": f"Gesamtbetrag ({total_suffix})",
                "value": result.calculated_total_amount,
                "unit": "€",
                "stored_label": "Gespeicherter Gesamtbetrag",
                "stored_value": obj.total_amount,
                "apply_url": apply_total_url
                if result.calculated_total_amount is not None
                else None,
            },
            {
                "label": f"Abschlag pro Monat ({installment_suffix})",
                "value": result.calculated_monthly_installment,
                "unit": "€",
                "stored_label": "Gespeicherter Abschlag",
                "stored_value": obj.monthly_installment,
                "apply_url": apply_installment_url
                if result.calculated_monthly_installment is not None
                else None,
            },
        ]

    def get_result_rows(self, obj: Supervision, result):
        from django.utils.formats import number_format

        i = result.input
        rows = []

        # Abrechnungsgrundlage — kontextabhängig, eine oder zwei Zeilen
        fee_forced = i.use_fee_agreement and result.pool_agreement is not None
        if result.is_pool_rate:
            rows.append(
                (
                    "Abrechnungsgrundlage (Poolvereinbarung)",
                    str(result.pool_agreement),
                    False,
                )
            )
        elif fee_forced:
            rows.append(
                (
                    "Abrechnungsgrundlage (Entgeltvereinbarung)",
                    str(result.fee_agreement) if result.fee_agreement else "—",
                    True,
                )
            )
            rows.append(
                ("Poolvereinbarung (nicht verwendet)", str(result.pool_agreement), True)
            )
        else:
            rows.append(
                (
                    "Abrechnungsgrundlage (Entgeltvereinbarung)",
                    str(result.fee_agreement) if result.fee_agreement else "—",
                    i.fee_agreement_override is not None,
                )
            )

        if not result.is_pool_rate and result.is_tandem:
            gross_total = result.calculated_total_amount * 2
            rows.append(
                (
                    "Bruttobetrag (100 %)",
                    f"{number_format(gross_total, decimal_pos=2, use_l10n=True)} €",
                    False,
                )
            )
            rows.append(("Tandemabzug", "50 %", True))

        bd = result.school_days_breakdown
        rows += [
            (
                "Schultage (effektiv)",
                bd["school_days"],
                i.school_days_override is not None,
            ),
            (
                "Urlaubstage (effektiv)",
                bd["vacation_days"],
                i.vacation_days_override is not None,
            ),
            (
                "Feiertage (effektiv)",
                bd["public_holidays"],
                i.public_holidays_override is not None,
            ),
            (
                "Gesamt Schultage (effektiv)",
                bd["total"],
                any(
                    v is not None
                    for v in (
                        i.school_days_override,
                        i.vacation_days_override,
                        i.public_holidays_override,
                    )
                ),
            ),
            (
                "Monate (rechnerisch)",
                number_format(obj.calculated_months, decimal_pos=1, use_l10n=True),
                False,
            ),
            (
                "Effektive Monate",
                number_format(result.months, decimal_pos=1, use_l10n=True),
                result.months != obj.calculated_months,
            ),
        ]
        return rows

    def run_calculation(self, obj: Supervision, overrides: dict):
        from .calculators import SupervisionCalculatorInput, run_supervision_calculation

        return run_supervision_calculation(
            SupervisionCalculatorInput(supervision=obj, **overrides)
        )

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        result = kwargs.get("result")
        obj = kwargs.get("obj") or self.get_object()
        if result is not None:
            i = result.input
            override_params = self.overrides_to_params(
                {
                    "months_override": i.months_override,
                    "school_days_override": i.school_days_override,
                    "vacation_days_override": i.vacation_days_override,
                    "public_holidays_override": i.public_holidays_override,
                    "fee_agreement_override": i.fee_agreement_override,
                    "use_fee_agreement": True if i.use_fee_agreement else None,
                }
            )
            ctx["docx_url"] = self.build_apply_url(
                obj, "admin:pedagogy_supervision_calculator_docx", override_params
            )
        return ctx


def build_supervision_docx(context: dict):
    from datetime import date

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from base.docx import BriefbogenDocument

    obj = context["supervision"]
    result = context["result"]
    bd = result.school_days_breakdown

    bdoc = BriefbogenDocument()

    # 1. Titel
    bdoc.add_title("Stunden- und Budgetkalkulation")
    bdoc.add_spacer()

    # 2. Metadaten-Tabelle (rahmenlos)
    meta_table = bdoc.add_borderless_table(rows=3, cols=2)
    meta_data = [
        ("Schulbegleitung für:", f"{obj.student.last_name}, {obj.student.first_name}"),
        ("Schule:", obj.school.name if obj.school else "–"),
        ("Zeitraum:", f"{obj.start_date:%d.%m.%Y} – {obj.end_date:%d.%m.%Y}"),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    bdoc.add_spacer()

    # 3. Schultage-Tabelle (Label | Zahl | "Tage")
    days_table = bdoc.add_grid_table(rows=4, cols=3)
    days_data = [
        ("Schultage:", bd.get("school_days", 0)),
        ("Urlaubsanspruch:", bd.get("vacation_days", 0)),
        ("Gesetzl. Feiertage:", bd.get("public_holidays", 0)),
        ("Insgesamt:", bd.get("total", 0)),
    ]
    for i, (label, count) in enumerate(days_data):
        row = days_table.rows[i]
        run = row.cells[0].paragraphs[0].add_run(label)
        run.bold = i == 3  # "Insgesamt" fett
        row.cells[1].text = str(count)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[2].text = "Tage"

    bdoc.add_spacer()

    # 4. Betreuungsbedarf
    bdoc.add_paragraph("Betreuungsbedarf lt. Stundenplan", bold=True)
    bdoc.add_paragraph(
        f"{context['weekly_hours_h']} Stunden {context['weekly_hours_min']} Minuten"
        f" = {context['weekly_minutes']} Minuten/Woche"
        f" = {context['daily_minutes']} Minuten/Tag"
    )
    bdoc.add_spacer()

    # 5. Berechnungsabschnitt
    total_school_days = bd.get("total", 0)
    bdoc.add_paragraph("Berechnung der Gesamtkosten und des Abschlages:", bold=True)
    total_hours_fmt = number_format(
        context["total_hours"], decimal_pos=2, use_l10n=True
    )
    bdoc.add_paragraph(
        f"{context['daily_minutes']} × {total_school_days} Schultage"
        f" = {context['total_minutes']} Minuten"
        f" = {total_hours_fmt} Stunden"
    )
    if context["hourly_rate"] is not None:
        total_amount_fmt = number_format(
            result.calculated_total_amount, decimal_pos=2, use_l10n=True
        )
        bdoc.add_paragraph(
            f"{context['hourly_rate']} € × {total_hours_fmt} Stunden"
            f" = {total_amount_fmt} €"
        )
    if result.calculated_total_amount is not None:
        total_amount_fmt = number_format(
            result.calculated_total_amount, decimal_pos=2, use_l10n=True
        )
        bdoc.add_paragraph(f"{total_amount_fmt} € Gesamtbetrag", bold=True)

    bdoc.add_spacer()

    # 6. Abschlag (optional)
    if result.calculated_monthly_installment is not None:
        monthly_fmt = number_format(
            result.calculated_monthly_installment, decimal_pos=2, use_l10n=True
        )
        bdoc.add_paragraph("Abschlag pro Monat:", bold=True)
        bdoc.add_paragraph(f"bezogen auf {result.months} Monate = {monthly_fmt} €")
        bdoc.add_spacer()

    # 7. Signaturblock
    bdoc.add_paragraph(f"Uslar, den {date.today():%d.%m.%Y}")
    bdoc.add_spacer()
    bdoc.add_paragraph("Schubert, Geschäftsführer")

    return bdoc


class SupervisionDocxView(BaseCalculatorView):
    title = "Betreuungsrechner"

    def get_form_class(self):
        from .forms import SupervisionCalculatorOverridesForm

        return SupervisionCalculatorOverridesForm

    def get(self, request, *args, **kwargs):
        from decimal import Decimal

        from .calculators import SupervisionCalculatorInput, run_supervision_calculation

        obj = self.get_object()
        overrides = self.parse_overrides(request.GET)
        result = run_supervision_calculation(
            SupervisionCalculatorInput(supervision=obj, **overrides)
        )

        # Vorberechnete Minutenwerte
        weekly_seconds = obj.weekly_hours.total_seconds() if obj.weekly_hours else 0
        weekly_hours_h = int(weekly_seconds // 3600)
        weekly_hours_min = int((weekly_seconds % 3600) // 60)
        weekly_minutes = int(weekly_seconds // 60)
        daily_minutes = weekly_minutes // 5
        total_school_days = result.school_days_breakdown["total"]
        total_minutes = daily_minutes * total_school_days
        total_hours = Decimal(total_minutes) / Decimal(60)

        if result.fee_agreement is not None:
            hourly_rate = (
                result.fee_agreement.price_tandem
                if result.is_tandem
                else result.fee_agreement.price_standard
            )
        else:
            hourly_rate = None

        context = {
            "supervision": obj,
            "result": result,
            "weekly_hours_h": weekly_hours_h,
            "weekly_hours_min": weekly_hours_min,
            "weekly_minutes": weekly_minutes,
            "daily_minutes": daily_minutes,
            "total_minutes": total_minutes,
            "total_hours": total_hours,
            "hourly_rate": hourly_rate,
        }

        bdoc = build_supervision_docx(context)
        filename = f"kalkulation_{obj.student.last_name}_{obj.student.first_name}.docx"
        return bdoc.to_response(filename)


class SupervisionBaseApplyView(BaseApplyView):
    title = "Betreuungsrechner"
    calculator_url_name = "admin:pedagogy_supervision_calculator"
    calculator_view_class = SupervisionCalculatorView


class SupervisionApplyTotalView(SupervisionBaseApplyView):
    def get_value(self, result):
        return result.calculated_total_amount

    def save_value(self, obj: Supervision, value) -> None:
        obj.total_amount = value
        obj.save(update_fields=["total_amount"])

    def error_message(self) -> str:
        return "Gesamtbetrag konnte nicht berechnet werden — keine Übernahme möglich."

    def success_message(self, formatted: str) -> str:
        return f"Gesamtbetrag übernommen: {formatted} €"


class SupervisionApplyInstallmentView(SupervisionBaseApplyView):
    def get_value(self, result):
        return result.calculated_monthly_installment

    def save_value(self, obj: Supervision, value) -> None:
        obj.monthly_installment = value
        obj.save(update_fields=["monthly_installment"])

    def error_message(self) -> str:
        return "Abschlag konnte nicht berechnet werden — keine Übernahme möglich."

    def success_message(self, formatted: str) -> str:
        return f"Abschlag übernommen: {formatted} €"


@admin.register(Student)
class StudentAdmin(BaseModelAdmin):
    list_display = ("full_name", "payer")
    search_fields = (
        "first_name",
        "last_name",
        "guardian_first_name",
        "guardian_last_name",
    )
    inlines = [AddressInline, PhoneInline, EmailInline]
    autocomplete_fields = ("payer",)
    fieldsets = [
        (
            "Stammdaten",
            {
                "fields": ["first_name", "middle_name", "last_name", "payer"],
                "classes": ["col-span-2"],
            },
        ),
        (
            "Erziehungsberechtigte:r",
            {
                "fields": ["guardian_first_name", "guardian_last_name"],
                "classes": ["col-span-1"],
            },
        ),
    ]

    def get_queryset(self, request):
        return super().get_queryset(request).select_related("payer")


class NewRequestListView(BaseUnionListView):
    """Kombinierte Übersichtsliste aus Betreuungen und Anträgen."""

    title = "Neuanträge"

    def get_filter_form_class(self):
        from .forms import NewRequestFilterForm

        return NewRequestFilterForm

    def get_queryset_a(self, request):
        return Supervision.objects.select_related("student", "school")

    def get_queryset_b(self, request):
        return Request.objects.select_related("student", "school")

    def get_columns(self) -> list[tuple[str, str | None]]:
        return [
            ("Typ", "_row_type"),
            ("Schüler:in", "student__last_name"),
            ("Schule", "school__name"),
            ("Von", "start_date"),
            ("Bis", "end_date"),
        ]

    def get_row(self, obj) -> list:
        if isinstance(obj, Supervision):
            typ = render_label("Betreuung", "info")
        else:
            typ = render_label("Antrag", "success")
        return [
            typ,
            str(obj.student),
            str(obj.school),
            obj.start_date,
            obj.end_date if obj.end_date else "–",
        ]


@admin.register(Supervision)
class SupervisionAdmin(BaseModelAdmin):
    list_display = (
        "student",
        "caretaker",
        "school",
        "start_date",
        "end_date",
        "weekly_hours",
        "is_prophylactic",
        "display_total_amount",
    )
    actions_detail = ["edit_action", "calculator_action"]
    list_filter_submit = True
    list_filter = ("school", ("start_date", RangeDateFilter))
    search_fields = ("student__first_name", "student__last_name")
    autocomplete_fields = ("student", "caretaker", "school")
    readonly_fields = ()
    fieldsets = [
        ("Schüler:in", {"fields": [("student", "is_prophylactic")]}),
        ("Betreuung", {"fields": ["caretaker", ("school", "class_name")]}),
        ("Zeitraum", {"fields": [("start_date", "end_date"), "weekly_hours"]}),
        ("Abrechnung", {"fields": [("total_amount", "monthly_installment")]}),
    ]

    def get_readonly_fields(self, request, obj=None):
        if obj is None:
            return ()
        return super().get_readonly_fields(request, obj)

    def get_search_results(self, request, queryset, search_term):
        queryset, use_distinct = super().get_search_results(
            request, queryset, search_term
        )
        return queryset.select_related("student", "caretaker"), use_distinct

    @action(
        description="Betreuungsrechner",
        url_path="calculator-action",
        permissions=["calculator_action"],
        variant=ActionVariant.DEFAULT,
    )
    def calculator_action(self, request, object_id: int):
        return redirect(
            reverse("admin:pedagogy_supervision_calculator", args=[object_id])
        )

    def has_calculator_action_permission(self, request, obj=None):
        return True

    @display(description="Gesamtbetrag", ordering="total_amount")
    def display_total_amount(self, obj: Supervision) -> str:
        if obj.total_amount is not None:
            formatted = number_format(obj.total_amount, decimal_pos=2, use_l10n=True)
            return format_html("{} €", formatted)
        url = reverse("admin:pedagogy_supervision_calculator", args=[obj.pk])
        return format_html('<a href="{}">Berechnen →</a>', url)

    def get_queryset(self, request):
        return (
            super()
            .get_queryset(request)
            .select_related("student", "student__payer", "caretaker", "school")
        )

    def get_urls(self):
        custom = [
            path(
                "<int:pk>/calculator/",
                self.admin_site.admin_view(
                    SupervisionCalculatorView.as_view(model_admin=self)
                ),
                name="pedagogy_supervision_calculator",
            ),
            path(
                "<int:pk>/calculator/apply-total/",
                self.admin_site.admin_view(
                    SupervisionApplyTotalView.as_view(model_admin=self)
                ),
                name="pedagogy_supervision_calculator_apply_total",
            ),
            path(
                "<int:pk>/calculator/apply-installment/",
                self.admin_site.admin_view(
                    SupervisionApplyInstallmentView.as_view(model_admin=self)
                ),
                name="pedagogy_supervision_calculator_apply_installment",
            ),
            path(
                "<int:pk>/calculator/docx/",
                self.admin_site.admin_view(
                    SupervisionDocxView.as_view(model_admin=self)
                ),
                name="pedagogy_supervision_calculator_docx",
            ),
        ]
        return custom + super().get_urls()


@admin.register(Request)
class RequestAdmin(BaseModelAdmin):
    list_display = (
        "student",
        "state",
        "start_date",
        "demand",
        "review_date",
        "decision_date",
        "approval_type",
    )
    list_filter = ("state", "approval_type")
    search_fields = ("student__first_name", "student__last_name", "notes")
    autocomplete_fields = ("student", "school")

    def get_queryset(self, request):
        return super().get_queryset(request).select_related("student", "school")


class SchoolContactInline(TabularInline):
    model = SchoolContact
    extra = 0
    hide_title = True


@admin.register(School)
class SchoolAdmin(BaseModelAdmin):
    search_fields = ("name",)
    inlines = [SchoolContactInline]


class TandemPairingForm(forms.ModelForm):
    class Meta:
        model = TandemPairing
        fields = "__all__"

    def clean(self):
        cleaned = super().clean() or {}
        a = cleaned.get("supervision_a")
        b = cleaned.get("supervision_b")
        if a and b:
            if a == b:
                raise forms.ValidationError(
                    "Eine Betreuung kann nicht mit sich selbst verbunden werden."
                )
            if a.caretaker_id != b.caretaker_id:
                raise forms.ValidationError(
                    "Beide Betreuungen müssen denselben Betreuer haben."
                )
            if a.school_id != b.school_id:
                raise forms.ValidationError(
                    "Beide Betreuungen müssen dieselbe Schule haben."
                )
            if a.weekly_hours != b.weekly_hours:
                raise forms.ValidationError(
                    "Beide Betreuungen müssen dieselbe Wochenstundenzahl haben."
                )
        return cleaned


@admin.register(TandemPairing)
class TandemPairingAdmin(BaseModelAdmin):
    form = TandemPairingForm
    list_display = [
        "__str__",
        "display_school",
        "display_caretaker",
        "display_weekly_hours",
        "display_caretaker_matches",
        "display_hours_match",
    ]
    readonly_fields = ["display_caretaker_matches", "display_hours_match"]
    autocomplete_fields = ["supervision_a", "supervision_b"]
    search_fields = [
        "supervision_a__student__first_name",
        "supervision_a__student__last_name",
        "supervision_b__student__first_name",
        "supervision_b__student__last_name",
    ]

    def get_queryset(self, request):
        return (
            super()
            .get_queryset(request)
            .select_related(
                "supervision_a__student",
                "supervision_a__school",
                "supervision_a__caretaker",
                "supervision_b__student",
                "supervision_b__caretaker",
            )
        )

    @display(description="Schule")
    def display_school(self, obj: TandemPairing) -> str:
        return obj.supervision_a.school.name

    @display(description="Betreuer:in")
    def display_caretaker(self, obj: TandemPairing) -> str:
        return obj.supervision_a.caretaker.full_name

    @display(description="Wochenstunden")
    def display_weekly_hours(self, obj: TandemPairing) -> str:
        return HourMinuteDurationField.format_std(obj.supervision_a.weekly_hours)

    @display(description="Betreuer stimmt überein", boolean=True)
    def display_caretaker_matches(self, obj: TandemPairing) -> bool:
        return obj.caretaker_matches

    @display(description="Stunden stimmen überein", boolean=True)
    def display_hours_match(self, obj: TandemPairing) -> bool:
        return obj.hours_match


_TANDEM_FILTER = Q(tandem_as_a__isnull=False) | Q(tandem_as_b__isnull=False)


class CostPayerFilter(SimpleListFilter):
    title = "Kostenträger"
    parameter_name = "kostentraeger"

    def lookups(self, request, model_admin):
        from finance.models import CostPayer

        qs = (
            CostPayer.objects.filter(students__supervisions__isnull=False)
            .distinct()
            .order_by("identifier")
        )
        return [(p.pk, str(p)) for p in qs]

    def queryset(self, request, queryset):
        if self.value():
            return queryset.filter(student__payer_id=self.value())
        return queryset


@admin.register(SchoolReport)
class SchoolReportAdmin(ReadOnlyAdminMixin, ListSummaryMixin, BaseModelAdmin):
    list_display = [
        "display_school",
        "display_student",
        "display_payer",
        "weekly_hours",
        "is_prophylactic",
        "display_is_tandem",
    ]
    list_filter = [
        ("school", AutocompleteSelectFilter),
        CostPayerFilter,
        ("start_date", RangeDateFilter),
    ]
    list_filter_submit = True
    list_display_links = None
    ordering = ["school__name", "student__last_name", "student__first_name"]
    search_fields = ["student__first_name", "student__last_name", "school__name"]

    def get_queryset(self, request):
        return (
            super()
            .get_queryset(request)
            .select_related("school", "student", "student__payer")
            .annotate(is_tandem=_TANDEM_FILTER)
        )

    def get_summary_sections(self, cl) -> list[list[dict]]:
        qs = cl.queryset
        agg = qs.aggregate(
            total_hours=Sum("weekly_hours"),
            hours_without_tandem=Sum("weekly_hours", filter=~_TANDEM_FILTER),
            hours_with_tandem=Sum("weekly_hours", filter=_TANDEM_FILTER),
            student_count=Count("student_id", distinct=True),
            school_count=Count("school_id", distinct=True),
            prophylactic_count=Count("id", filter=Q(is_prophylactic=True)),
            tandem_count=Count("id", filter=_TANDEM_FILTER),
        )
        fmt = HourMinuteDurationField.format_std
        return [
            [
                {
                    "label": "Gesamtstunden",
                    "value": fmt(agg["total_hours"]) if agg["total_hours"] else "–",
                },
                {
                    "label": "Std. ohne Tandem",
                    "value": fmt(agg["hours_without_tandem"])
                    if agg["hours_without_tandem"]
                    else "–",
                },
                {
                    "label": "Std. mit Tandem",
                    "value": fmt(agg["hours_with_tandem"])
                    if agg["hours_with_tandem"]
                    else "–",
                },
            ],
            [
                {"label": "Schüler:innen", "value": agg["student_count"] or 0},
                {"label": "Schulen", "value": agg["school_count"] or 0},
                {"label": "Prophylaktisch", "value": agg["prophylactic_count"] or 0},
                {"label": "Tandem", "value": agg["tandem_count"] or 0},
            ],
        ]

    @display(description="Schule", ordering="school__name")
    def display_school(self, obj: SchoolReport) -> str:
        return obj.school.name

    @display(description="Schüler:in", ordering="student__last_name")
    def display_student(self, obj: SchoolReport) -> str:
        return obj.student.full_name

    @display(description="Kostenträger", ordering="student__payer__identifier")
    def display_payer(self, obj: SchoolReport) -> str:
        return str(obj.student.payer)

    @display(description="Tandem", boolean=True)
    def display_is_tandem(self, obj: SchoolReport) -> bool:
        return bool(obj.is_tandem)  # type: ignore[attr-defined]
